"""
DOCX Tool - Microsoft Word document processing.

Provides capabilities for:
- Reading DOCX documents
- Extracting text and structure
- Creating and modifying documents
- Table extraction
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from tools.base import Tool, ToolParameter, ToolResult, register_tool

logger = logging.getLogger("titan.tools.documents.docx")


class DOCXTool(Tool):
    """
    Microsoft Word document tool.

    Allows agents to read and manipulate DOCX files.
    """

    @property
    def name(self) -> str:
        return "docx"

    @property
    def description(self) -> str:
        return (
            "Microsoft Word (DOCX) document tool. "
            "Actions: 'read' (extract text), 'structure' (get document structure), "
            "'tables' (extract tables), 'create' (create new document)."
        )

    @property
    def parameters(self) -> list[ToolParameter]:
        return [
            ToolParameter(
                name="action",
                type="string",
                description="Action: 'read', 'structure', 'tables', 'create'",
                required=True,
                enum=["read", "structure", "tables", "create"],
            ),
            ToolParameter(
                name="path",
                type="string",
                description="Path to DOCX file (for read actions) or output path (for create)",
                required=True,
            ),
            ToolParameter(
                name="content",
                type="string",
                description="Content to write (for 'create' action)",
                required=False,
            ),
            ToolParameter(
                name="max_chars",
                type="integer",
                description="Maximum characters to return (default: 50000)",
                required=False,
            ),
        ]

    async def execute(self, **kwargs: Any) -> ToolResult:
        try:
            from docx import Document
        except ImportError:
            return ToolResult(
                success=False,
                output=None,
                error="python-docx is required. Install with: pip install 'agentic-titan[documents]'",
            )

        action = kwargs.get("action", "")
        path = kwargs.get("path", "")
        max_chars = kwargs.get("max_chars", 50000)

        if not path:
            return ToolResult(
                success=False,
                output=None,
                error="'path' parameter is required",
            )

        try:
            if action == "read":
                doc = Document(path)
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)

                text = "\n\n".join(paragraphs)
                if len(text) > max_chars:
                    text = text[:max_chars] + "\n\n[Truncated]"

                return ToolResult(
                    success=True,
                    output={
                        "text": text,
                        "paragraph_count": len(paragraphs),
                        "total_chars": sum(len(p) for p in paragraphs),
                    },
                )

            elif action == "structure":
                doc = Document(path)
                structure = []

                for para in doc.paragraphs:
                    if para.text.strip():
                        style = para.style.name if para.style else "Normal"
                        structure.append({
                            "type": "paragraph",
                            "style": style,
                            "text_preview": para.text[:100] + "..." if len(para.text) > 100 else para.text,
                        })

                for table in doc.tables:
                    structure.append({
                        "type": "table",
                        "rows": len(table.rows),
                        "cols": len(table.columns) if table.rows else 0,
                    })

                return ToolResult(
                    success=True,
                    output={
                        "elements": len(structure),
                        "structure": structure[:50],  # Limit for large docs
                    },
                )

            elif action == "tables":
                doc = Document(path)
                tables_data = []

                for i, table in enumerate(doc.tables):
                    rows_data = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows_data.append(cells)

                    tables_data.append({
                        "table_index": i,
                        "rows": len(table.rows),
                        "data": rows_data[:20],  # Limit rows
                    })

                return ToolResult(
                    success=True,
                    output={
                        "table_count": len(tables_data),
                        "tables": tables_data,
                    },
                )

            elif action == "create":
                content = kwargs.get("content", "")
                if not content:
                    return ToolResult(
                        success=False,
                        output=None,
                        error="'content' parameter required for create action",
                    )

                doc = Document()

                # Split content into paragraphs
                paragraphs = content.split("\n\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        # Check for heading markers
                        if para_text.startswith("# "):
                            doc.add_heading(para_text[2:], level=1)
                        elif para_text.startswith("## "):
                            doc.add_heading(para_text[3:], level=2)
                        elif para_text.startswith("### "):
                            doc.add_heading(para_text[4:], level=3)
                        else:
                            doc.add_paragraph(para_text.strip())

                doc.save(path)

                return ToolResult(
                    success=True,
                    output={
                        "created": True,
                        "path": path,
                        "paragraph_count": len(paragraphs),
                    },
                )

            else:
                return ToolResult(
                    success=False,
                    output=None,
                    error=f"Unknown action: {action}",
                )

        except FileNotFoundError:
            return ToolResult(
                success=False,
                output=None,
                error=f"File not found: {path}",
            )
        except Exception as e:
            logger.exception(f"DOCX tool error: {e}")
            return ToolResult(
                success=False,
                output=None,
                error=f"Failed to process DOCX: {e}",
            )


# Register the tool
docx_tool = DOCXTool()
register_tool(docx_tool)
